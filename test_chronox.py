# test_chronox.py
import unittest
import os
import pandas as pd
from docx import Document
from chronox import DateExtractor # Updated import

class TestDateExtractor(unittest.TestCase):
    """
    Unit tests for the DateExtractor class.
    """

    @classmethod
    def setUpClass(cls):
        """
        Set up test resources once for all tests.
        Creates dummy .docx files for testing.
        """
        cls.test_dir = "test_docs"
        os.makedirs(cls.test_dir, exist_ok=True)

        # Document with various date formats
        cls.doc_path_1 = os.path.join(cls.test_dir, "test_document_1.docx")
        document1 = Document()
        document1.add_paragraph("This is a test document with dates.")
        document1.add_paragraph("The project started on 01/01/2023 and ended on 31-12-2023.")
        document1.add_paragraph("Another important date is 2024-05-15.")
        document1.add_paragraph("The meeting was held on January 15, 2024.")
        document1.add_paragraph("The deadline is 1 February 2025.")
        document1.add_paragraph("The contract was signed in March 2023.")
        document1.save(cls.doc_path_1)

        # Document with no dates
        cls.doc_path_no_dates = os.path.join(cls.test_dir, "test_document_no_dates.docx")
        document_no_dates = Document()
        document_no_dates.add_paragraph("This document contains no specific dates.")
        document_no_dates.add_paragraph("It's just plain text.")
        document_no_dates.save(cls.doc_path_no_dates)

        # Empty document
        cls.doc_path_empty = os.path.join(cls.test_dir, "test_document_empty.docx")
        document_empty = Document()
        document_empty.save(cls.doc_path_empty)

        # Document with multiple dates in one sentence
        cls.doc_path_multi_dates = os.path.join(cls.test_dir, "test_document_multi_dates.docx")
        document_multi_dates = Document()
        document_multi_dates.add_paragraph("Events occurred on 01/01/2023 and 02/02/2024.")
        document_multi_dates.save(cls.doc_path_multi_dates)

    @classmethod
    def tearDownClass(cls):
        """
        Clean up test resources after all tests are run.
        Removes dummy .docx and output files.
        """
        # Clean up test documents
        for f in os.listdir(cls.test_dir):
            os.remove(os.path.join(cls.test_dir, f))
        os.rmdir(cls.test_dir)

        # Clean up any generated output files
        if os.path.exists("extracted_events.csv"):
            os.remove("extracted_events.csv")
        if os.path.exists("output.csv"):
            os.remove("output.csv")
        if os.path.exists("output.xlsx"):
            os.remove("output.xlsx")
        if os.path.exists(os.path.join(cls.test_dir, "empty_output.csv")): # Ensure this is also cleaned
            os.remove(os.path.join(cls.test_dir, "empty_output.csv"))
        if os.path.exists(os.path.join(cls.test_dir, "test.txt")): # Ensure this is also cleaned
            os.remove(os.path.join(cls.test_dir, "test.txt"))


    def setUp(self):
        """
        Set up for each test method.
        Initializes a new DateExtractor instance for each test.
        """
        self.extractor = DateExtractor()

    def test_extract_dates_from_docx_valid_document(self):
        """
        Test date extraction from a document with various valid date formats.
        """
        df = self.extractor.extract_dates_from_docx(self.doc_path_1)
        self.assertIsInstance(df, pd.DataFrame)
        self.assertFalse(df.empty)
        self.assertIn('Date Found', df.columns)
        self.assertIn('Event Description', df.columns)

        expected_dates = [
            "01/01/2023", "31-12-2023", "2024-05-15",
            "January 15, 2024", "1 February 2025", "March 2023"
        ]
        # Check if all expected dates are present in the 'Date Found' column
        for date_str in expected_dates:
            self.assertIn(date_str, df['Date Found'].tolist())

        # Check for correct event descriptions (sentences)
        self.assertIn("The project started on 01/01/2023 and ended on 31-12-2023.", df['Event Description'].tolist())
        self.assertIn("Another important date is 2024-05-15.", df['Event Description'].tolist())
        self.assertIn("The meeting was held on January 15, 2024.", df['Event Description'].tolist())
        self.assertIn("The deadline is 1 February 2025.", df['Event Description'].tolist())
        self.assertIn("The contract was signed in March 2023.", df['Event Description'].tolist())


    def test_extract_dates_from_docx_no_dates(self):
        """
        Test date extraction from a document with no dates.
        Should return an empty DataFrame with correct columns.
        """
        df = self.extractor.extract_dates_from_docx(self.doc_path_no_dates)
        self.assertIsInstance(df, pd.DataFrame)
        self.assertTrue(df.empty)
        self.assertIn('Date Found', df.columns) # More robust check
        self.assertIn('Event Description', df.columns) # More robust check

    def test_extract_dates_from_docx_empty_document(self):
        """
        Test date extraction from an empty document.
        Should return an empty DataFrame with correct columns.
        """
        df = self.extractor.extract_dates_from_docx(self.doc_path_empty)
        self.assertIsInstance(df, pd.DataFrame)
        self.assertTrue(df.empty)
        self.assertIn('Date Found', df.columns) # More robust check
        self.assertIn('Event Description', df.columns) # More robust check

    def test_extract_dates_from_docx_non_existent_file(self):
        """
        Test date extraction with a non-existent file path.
        Should return an empty DataFrame with correct columns.
        """
        df = self.extractor.extract_dates_from_docx("non_existent_file.docx")
        self.assertIsInstance(df, pd.DataFrame)
        self.assertTrue(df.empty)
        self.assertIn('Date Found', df.columns) # More robust check
        self.assertIn('Event Description', df.columns) # More robust check

    def test_extract_dates_from_docx_wrong_file_type(self):
        """
        Test date extraction with a file that is not a .docx.
        Should return an empty DataFrame with correct columns.
        """
        # Create a dummy text file
        with open(os.path.join(self.test_dir, "test.txt"), "w") as f:
            f.write("This is a text file.")
        df = self.extractor.extract_dates_from_docx(os.path.join(self.test_dir, "test.txt"))
        self.assertIsInstance(df, pd.DataFrame)
        self.assertTrue(df.empty)
        self.assertIn('Date Found', df.columns) # More robust check
        self.assertIn('Event Description', df.columns) # More robust check

    def test_extract_dates_from_docx_multi_dates_in_sentence(self):
        """
        Test extraction when multiple dates are in the same sentence.
        Each date should result in a separate entry with the same event description.
        """
        df = self.extractor.extract_dates_from_docx(self.doc_path_multi_dates)
        self.assertIsInstance(df, pd.DataFrame)
        self.assertFalse(df.empty)
        self.assertEqual(len(df), 2) # Expect two entries
        self.assertIn("01/01/2023", df['Date Found'].tolist())
        self.assertIn("02/02/2024", df['Date Found'].tolist())
        # Both should have the same event description
        self.assertEqual(df['Event Description'].iloc[0], "Events occurred on 01/01/2023 and 02/02/2024.")
        self.assertEqual(df['Event Description'].iloc[1], "Events occurred on 01/01/2023 and 02/02/2024.")


    def test_save_to_spreadsheet_csv(self):
        """
        Test saving DataFrame to a CSV file.
        """
        test_df = pd.DataFrame({
            'Date Found': ['01/01/2023', '02/02/2024'],
            'Event Description': ['Event A', 'Event B']
        })
        output_csv_path = os.path.join(self.test_dir, "output.csv")
        self.extractor.save_to_spreadsheet(test_df, output_csv_path)
        self.assertTrue(os.path.exists(output_csv_path))

        # Verify content
        read_df = pd.read_csv(output_csv_path)
        pd.testing.assert_frame_equal(test_df, read_df)

    def test_save_to_spreadsheet_xlsx(self):
        """
        Test saving DataFrame to an XLSX file.
        """
        test_df = pd.DataFrame({
            'Date Found': ['01/01/2023', '02/02/2024'],
            'Event Description': ['Event A', 'Event B']
        })
        output_xlsx_path = os.path.join(self.test_dir, "output.xlsx")
        self.extractor.save_to_spreadsheet(test_df, output_xlsx_path)
        self.assertTrue(os.path.exists(output_xlsx_path))

        # Verify content
        read_df = pd.read_excel(output_xlsx_path)
        pd.testing.assert_frame_equal(test_df, read_df)

    def test_save_to_spreadsheet_empty_df(self):
        """
        Test saving an empty DataFrame.
        Should create a file with headers, and the DataFrame read back should be empty with correct columns.
        """
        empty_df = pd.DataFrame(columns=['Date Found', 'Event Description'])
        output_csv_path = os.path.join(self.test_dir, "empty_output.csv")
        self.extractor.save_to_spreadsheet(empty_df, output_csv_path)
        self.assertTrue(os.path.exists(output_csv_path))

        # Verify content: file exists, is readable, is empty, and has the expected columns
        read_df = pd.read_csv(output_csv_path)
        self.assertTrue(read_df.empty)
        self.assertIn('Date Found', read_df.columns) # More robust check
        self.assertIn('Event Description', read_df.columns) # More robust check

    def test_save_to_spreadsheet_unsupported_format(self):
        """
        Test saving to an unsupported file format.
        Should not create a file.
        """
        test_df = pd.DataFrame({
            'Date Found': ['01/01/2023'],
            'Event Description': ['Event A']
        })
        output_txt_path = os.path.join(self.test_dir, "output.txt")
        self.extractor.save_to_spreadsheet(test_df, output_txt_path)
        self.assertFalse(os.path.exists(output_txt_path))


if __name__ == '__main__':
    unittest.main(argv=['first-arg-is-ignored'], exit=False)

